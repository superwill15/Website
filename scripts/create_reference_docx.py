"""
Create AssetStage branded reference DOCX template for pandoc conversion.
Uses python-docx to create a template with:
- Branded heading styles (Navy #1e3c72)
- Header with logo
- Footer with contact info and page numbers
- Professional table styling
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# Brand colors
BRAND_NAVY = RGBColor(0x1e, 0x3c, 0x72)  # #1e3c72 - Primary Navy for headers
BRAND_BLUE = RGBColor(0x34, 0x98, 0xdb)  # #3498db - Accent Blue for links
BRAND_GREEN = RGBColor(0x27, 0xae, 0x60)  # #27ae60 - Accent Green for CTAs
TEXT_DARK = RGBColor(0x2c, 0x3e, 0x50)   # #2c3e50 - Body text

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    """Add page number field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_header_logo(header, logo_path):
    """Add logo to the header."""
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    if os.path.exists(logo_path):
        run.add_picture(logo_path, width=Inches(1.8))

def create_reference_docx(output_path, logo_path):
    """Create the branded reference DOCX template."""
    doc = Document()

    # Set up page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure styles
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(28)
    title_font.color.rgb = BRAND_NAVY
    title_font.bold = True
    title_style.paragraph_format.space_after = Pt(6)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle style
    subtitle_style = styles['Subtitle']
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Calibri'
    subtitle_font.size = Pt(14)
    subtitle_font.color.rgb = TEXT_DARK
    subtitle_font.italic = True
    subtitle_style.paragraph_format.space_after = Pt(24)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Calibri'
    h1_font.size = Pt(18)
    h1_font.color.rgb = BRAND_NAVY
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(8)

    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.color.rgb = BRAND_NAVY
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(6)

    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.color.rgb = BRAND_NAVY
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(4)

    # Heading 4 style
    h4_style = styles['Heading 4']
    h4_font = h4_style.font
    h4_font.name = 'Calibri'
    h4_font.size = Pt(11)
    h4_font.color.rgb = BRAND_NAVY
    h4_font.bold = True
    h4_font.italic = False
    h4_style.paragraph_format.space_before = Pt(10)
    h4_style.paragraph_format.space_after = Pt(4)

    # Normal text style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = TEXT_DARK
    normal_style.paragraph_format.space_after = Pt(8)
    normal_style.paragraph_format.line_spacing = 1.15

    # Code/Source Code style for code blocks
    try:
        code_style = styles.add_style('Source Code', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code_style = styles['Source Code']
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_font.color.rgb = TEXT_DARK
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Hyperlink character style
    try:
        hyperlink_style = styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
        hyperlink_font = hyperlink_style.font
        hyperlink_font.color.rgb = BRAND_BLUE
        hyperlink_font.underline = True
    except ValueError:
        pass  # Style already exists

    # Set up header with logo
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Add logo to header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(logo_path):
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.5))

    # Add a thin line under the header
    header_para.paragraph_format.space_after = Pt(6)

    # Set up footer
    footer = section.footer
    footer.is_linked_to_previous = False

    # Footer paragraph with contact info and page number
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add contact info
    run = footer_para.add_run("www.assetstage.io  |  sales@assetstage.io  |  Page ")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_DARK

    # Add page number
    add_page_number(footer_para)

    # Add sample content to show the template (will be replaced by pandoc)
    doc.add_paragraph("Sample Title", style='Title')
    doc.add_paragraph("Sample Subtitle", style='Subtitle')
    doc.add_paragraph("Heading 1 Example", style='Heading 1')
    doc.add_paragraph("This is normal body text with the brand styling.", style='Normal')
    doc.add_paragraph("Heading 2 Example", style='Heading 2')
    doc.add_paragraph("More body text to demonstrate the formatting.", style='Normal')
    doc.add_paragraph("Heading 3 Example", style='Heading 3')
    doc.add_paragraph("Additional content here.", style='Normal')

    # Add a sample table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, text in enumerate(['Column 1', 'Column 2', 'Column 3']):
        cell = header_cells[i]
        cell.text = text
        # Style the header cell
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Set background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1e3c72')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx in range(1, 3):
        for col_idx in range(3):
            table.rows[row_idx].cells[col_idx].text = f'Data {row_idx},{col_idx}'

    # Save the document
    doc.save(output_path)
    print(f"Reference DOCX created: {output_path}")

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)

    output_path = os.path.join(project_root, "content", "resources", "assetstage-reference.docx")
    logo_path = os.path.join(project_root, "public", "logo-primary.png")

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    create_reference_docx(output_path, logo_path)
